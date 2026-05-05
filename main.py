from bs4 import BeautifulSoup
import requests 
from docx import Document
from docx.shared import Pt
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
    

def set_styles(font_name,font_size,document_styles,document):
    style = document.styles[document_styles]
    style.font.name = font_name
    style.font.size = Pt(int(font_size))



link = input('введите ссылку:  ') 
while "wikipedia.org" not in link:
    print("ссылка не в википедию. Введите еще раз: ")
    link = input('введите ссылку: ') 

headers = {"user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"}
response = requests.get(url = link, headers = headers )
response.raise_for_status()

soup = BeautifulSoup(response.text, features="html.parser")
document = Document()




while True:
    style = input("Введите стиль: ")
    name = input("Введите имя шрифта: ")
    size = input("введите размер шрифта: ")
    try:
        set_styles(font_name=name ,font_size=size ,document_styles=style ,document = document)
        break
    except KeyError:
        print("стиля с таким названием не существует")
    except ValueError:
        print("размера такого не существует")



clear_report = soup.find('div', class_="vector-body")
tags = clear_report.find_all(['p','h2','h1','img'])

reveersed_tags =  list(reversed(tags))

for index,i  in enumerate(reveersed_tags):
        if "<p"  in str(i):    
            break
    
without_empty_sections_tags = reveersed_tags[index:len(reveersed_tags)]
without_empty_sections_tags = list(reversed(without_empty_sections_tags)) 

mask = r"\[.*?\]"
for tag in without_empty_sections_tags:
    clear_beautiful_text=(re.sub(mask, 
                r'',
                tag.text))
    

    if "<p" in str(tag):
        paragraph = document.add_paragraph(clear_beautiful_text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
    if "<h2" in str(tag):
        head = document.add_heading(clear_beautiful_text, level=2)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    if "<h1" in str(tag):
        head = document.add_heading(clear_beautiful_text,level=1)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    if "<img" in str(tag):
        img_link = tag["src"]
        response = requests.get(url = 'https:'+ img_link, headers = headers)
        response.raise_for_status()
        picture = response.content
        with open("myfile.png", "wb",) as my_file:
            my_file.write(picture)
        try:
            document_pucture = document.add_picture('myfile.png')
            picture_paragraph = document.paragraphs[-1]
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
             print(f"картинка {'https:'+ link} не добавлена")

             
while True:
    try:
        document.save('result.docx')
        break
    except PermissionError:
        input("закройте word-файл и нажминьте enter")
        
# os.startfile('result.docx') 
